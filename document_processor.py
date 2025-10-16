"""
Document processing and content analysis for search indexing
"""

import os
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import re

# Document processing libraries
import PyPDF2
import docx
import pandas as pd
from bs4 import BeautifulSoup

try:
    from .config import DocumentMetadata, search_config
except ImportError:
    from config import DocumentMetadata, search_config
try:
    from .models import DocumentChunk
except ImportError:
    from models import DocumentChunk

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes various document types for search indexing"""

    def __init__(self):
        self.supported_extensions = search_config.supported_extensions

    def can_process(self, file_path: str) -> bool:
        """Check if document type is supported"""
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_extensions

    def extract_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content and metadata from document

        Returns:
            Tuple of (content_text, metadata_dict)
        """
        if not self.can_process(file_path):
            raise ValueError(f"Unsupported file type: {file_path}")

        extension = Path(file_path).suffix.lower()

        try:
            if extension == '.pdf':
                return self._extract_pdf_content(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_docx_content(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._extract_excel_content(file_path)
            elif extension in ['.html', '.htm']:
                return self._extract_html_content(file_path)
            elif extension == '.txt':
                return self._extract_text_content(file_path)
            elif extension == '.md':
                return self._extract_markdown_content(file_path)
            elif extension == '.rtf':
                return self._extract_rtf_content(file_path)
            else:
                return self._extract_text_content(file_path)

        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {str(e)}")
            return "", {}

    def _extract_pdf_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF files"""
        content = ""
        metadata = {}

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                        'modification_date': pdf_reader.metadata.get('/ModDate', ''),
                    }

                # Extract text content
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"

        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")

        return content.strip(), metadata

    def _extract_docx_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from Word documents"""
        content = ""
        metadata = {}

        try:
            doc = docx.Document(file_path)

            # Extract document properties
            if doc.core_properties:
                metadata = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'keywords': doc.core_properties.keywords or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                }

            # Extract text content
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " | "
                    content += row_text.rstrip(" | ") + "\n"

        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")

        return content.strip(), metadata

    def _extract_excel_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from Excel files"""
        content = ""
        metadata = {}

        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)

            metadata = {
                'sheet_names': excel_file.sheet_names,
                'num_sheets': len(excel_file.sheet_names),
            }

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Add sheet name as header
                content += f"\n=== {sheet_name} ===\n"

                # Convert dataframe to text
                content += df.to_string(index=False) + "\n"

        except Exception as e:
            logger.error(f"Error reading Excel {file_path}: {str(e)}")

        return content.strip(), metadata

    def _extract_html_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from HTML files"""
        content = ""
        metadata = {}

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')

                # Extract metadata from meta tags
                for meta in soup.find_all('meta'):
                    name = meta.get('name', '').lower()
                    property_name = meta.get('property', '').lower()
                    content_value = meta.get('content', '')

                    if name == 'description':
                        metadata['description'] = content_value
                    elif name == 'keywords':
                        metadata['keywords'] = content_value
                    elif name == 'author':
                        metadata['author'] = content_value
                    elif property_name == 'og:title':
                        metadata['title'] = content_value

                # Extract title
                if soup.title:
                    metadata['title'] = soup.title.string or ''

                # Extract text content (remove scripts and styles)
                for script in soup(["script", "style"]):
                    script.decompose()

                content = soup.get_text(separator="\n")

        except Exception as e:
            logger.error(f"Error reading HTML {file_path}: {str(e)}")

        return content.strip(), metadata

    def _extract_text_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content.strip(), {}
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return "", {}

    def _extract_markdown_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from Markdown files"""
        return self._extract_text_content(file_path)

    def _extract_rtf_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from RTF files"""
        # For now, treat as text file
        # Could be enhanced with python-rtf library if needed
        return self._extract_text_content(file_path)

    def extract_document_metadata(self, file_path: str, doc_id: str) -> DocumentMetadata:
        """Extract comprehensive document metadata"""
        path_obj = Path(file_path)
        stat = path_obj.stat()

        # Basic file metadata
        metadata = DocumentMetadata(
            doc_id=doc_id,
            file_path=str(path_obj.relative_to(Path.cwd())),
            absolute_path=str(path_obj.absolute()),
            file_name=path_obj.name,
            file_extension=path_obj.suffix.lower(),
            file_size_bytes=stat.st_size,
            file_size_human=self._format_file_size(stat.st_size),
            date_created=datetime.fromtimestamp(stat.st_ctime).isoformat(),
            date_modified=datetime.fromtimestamp(stat.st_mtime).isoformat(),
            date_indexed=datetime.now().isoformat()
        )

        # Try to extract additional metadata from content
        try:
            content, content_metadata = self.extract_content(file_path)

            # Extract enhanced metadata from content
            metadata.title = self._extract_title(content, content_metadata)
            metadata.author = content_metadata.get('author', '')
            metadata.content_type = self._determine_content_type(content, path_obj)
            metadata.category = self._extract_category(content, path_obj)
            metadata.document_type = self._determine_document_type(content, path_obj)
            metadata.version = self._extract_version(content)

        except Exception as e:
            logger.warning(f"Could not extract enhanced metadata from {file_path}: {str(e)}")

        return metadata

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"

    def _extract_title(self, content: str, metadata: Dict[str, Any]) -> str:
        """Extract document title from content"""
        # Try metadata first
        if metadata.get('title'):
            return metadata['title']

        # Try to extract from first line
        lines = content.strip().split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line) > 10 and len(line) < 200:
                # Skip version control lines
                if not re.match(r'^version\s+\d+', line.lower()):
                    return line

        # Fallback to filename without extension
        return Path(metadata.get('file_path', '')).stem if 'file_path' in metadata else ''

    def _determine_content_type(self, content: str, path_obj: Path) -> str:
        """Determine content type based on content and path"""
        content_lower = content.lower()
        path_str = str(path_obj).lower()

        if 'invoice' in content_lower or 'invoice' in path_str:
            return 'Invoice Processing'
        elif 'payment' in content_lower or 'payment' in path_str:
            return 'Payment Processing'
        elif 'vendor' in content_lower or 'vendor' in path_str:
            return 'Vendor Management'
        elif 'helpdesk' in content_lower or 'helpdesk' in path_str:
            return 'Helpdesk Procedures'
        elif 'escalation' in content_lower or 'escalation' in path_str:
            return 'Escalation Procedures'
        elif 'period close' in content_lower or 'period close' in path_str:
            return 'Period Close'
        else:
            return 'General Documentation'

    def _extract_category(self, content: str, path_obj: Path) -> str:
        """Extract category from path structure"""
        path_parts = path_obj.parts

        # Look for category indicators in path
        for part in path_parts:
            if 'approved for use' in part.lower():
                return 'Approved Procedures'
            elif 'draft' in part.lower():
                return 'Draft Documents'
            elif 'templates' in part.lower():
                return 'Templates'
            elif 'guides' in part.lower():
                return 'Guides'

        return 'General'

    def _determine_document_type(self, content: str, path_obj: Path) -> str:
        """Determine document type based on content and extension"""
        extension = path_obj.suffix.lower()

        if extension in ['.pdf', '.docx', '.doc']:
            # Analyze content for document type
            content_lower = content.lower()

            if any(keyword in content_lower for keyword in ['procedure', 'process', 'steps to', 'how to']):
                return 'Procedure'
            elif any(keyword in content_lower for keyword in ['checklist', 'verification', 'validation']):
                return 'Checklist'
            elif any(keyword in content_lower for keyword in ['table', 'matrix', 'reference']):
                return 'Reference'
            elif any(keyword in content_lower for keyword in ['guide', 'manual', 'handbook']):
                return 'Guide'
            else:
                return 'Document'

        elif extension in ['.xlsx', '.xls']:
            return 'Spreadsheet'
        elif extension in ['.html', '.htm']:
            return 'Web Document'
        else:
            return 'Text Document'

    def _extract_version(self, content: str) -> str:
        """Extract version information from content"""
        # Look for version patterns
        version_patterns = [
            r'version[\s:]+(\d+(?:\.\d+)*)',
            r'v\d+(?:\.\d+)*',
            r'ver[\s\.]+(\d+(?:\.\d+)*)',
        ]

        for pattern in version_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1) or match.group(0)

        return ''

    def chunk_content(self, content: str, doc_id: str, chunk_size: int = None, overlap: int = None) -> List[DocumentChunk]:
        """Split document content into searchable chunks"""
        if chunk_size is None:
            chunk_size = search_config.chunk_size
        if overlap is None:
            overlap = search_config.chunk_overlap

        if not content:
            return []

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(content):
            end = start + chunk_size

            # Try to end at a sentence or word boundary
            if end < len(content):
                # Look for sentence endings
                sentence_end = content.rfind('.', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
                else:
                    # Look for word boundaries
                    word_end = content.rfind(' ', start, end)
                    if word_end > start + chunk_size // 2:
                        end = word_end

            chunk_text = content[start:end].strip()
            if chunk_text:
                chunk = DocumentChunk(
                    chunk_id=f"{doc_id}_chunk_{chunk_index}",
                    doc_id=doc_id,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                    token_count=len(chunk_text.split())
                )
                chunks.append(chunk)

            # Move start position with overlap
            start = end - overlap if end < len(content) else end
            chunk_index += 1

            # Prevent infinite loop
            if start >= end:
                break

        return chunks


class ContentAnalyzer:
    """Analyzes document content for enhanced search capabilities"""

    def __init__(self):
        self.stop_words = self._load_stop_words()

    def _load_stop_words(self) -> set:
        """Load common stop words for text analysis"""
        # Basic English stop words
        return {
            'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from',
            'has', 'he', 'in', 'is', 'it', 'its', 'of', 'on', 'that', 'the',
            'to', 'was', 'will', 'with', 'have', 'had', 'has', 'do', 'does',
            'did', 'but', 'or', 'not', 'no', 'yes', 'this', 'these', 'those'
        }

    def extract_keywords(self, content: str, max_keywords: int = 20) -> List[Tuple[str, float]]:
        """Extract important keywords and their weights from content"""
        if not content:
            return []

        # Simple keyword extraction based on frequency
        words = re.findall(r'\b\w+\b', content.lower())
        word_count = {}

        for word in words:
            if len(word) > 3 and word not in self.stop_words:
                word_count[word] = word_count.get(word, 0) + 1

        # Calculate weights based on frequency and position
        total_words = len(words)
        keywords = []

        for word, count in sorted(word_count.items(), key=lambda x: x[1], reverse=True):
            if len(keywords) >= max_keywords:
                break

            # Boost score for words appearing in title-like positions
            boost = 1.0
            if word in content[:500].lower():  # First 500 chars
                boost = 1.5

            weight = (count / total_words) * boost
            keywords.append((word, weight))

        return keywords

    def extract_entities(self, content: str) -> Dict[str, List[str]]:
        """Extract named entities from content"""
        entities = {
            'systems': [],
            'departments': [],
            'processes': [],
            'terms': []
        }

        content_lower = content.lower()

        # System entities
        systems = [
            'oracle cloud', 'oracle ebs', 'sap', 'servicenow', 'maximo',
            'ariba', 'lams', 'entrac', 'openbill', 'peoplesoft'
        ]

        for system in systems:
            if system in content_lower:
                entities['systems'].append(system.title())

        # Department entities
        departments = [
            'accounts payable', 'vendor master', 'supply chain management',
            'helpdesk', 'finance', 'procurement', 'legal', 'compliance'
        ]

        for dept in departments:
            if dept in content_lower:
                entities['departments'].append(dept.title())

        # Process entities
        processes = [
            'invoice processing', 'payment processing', 'vendor management',
            'period close', 'month end close', 'reconciliation', 'escalation'
        ]

        for process in processes:
            if process in content_lower:
                entities['processes'].append(process.title())

        return entities

    def analyze_sentiment(self, content: str) -> Dict[str, float]:
        """Basic sentiment analysis for document tone"""
        # Simple sentiment analysis based on positive/negative words
        positive_words = {'good', 'excellent', 'effective', 'efficient', 'approved', 'success'}
        negative_words = {'error', 'issue', 'problem', 'failed', 'rejected', 'cancelled', 'hold'}

        words = set(re.findall(r'\b\w+\b', content.lower()))

        positive_count = len(words.intersection(positive_words))
        negative_count = len(words.intersection(negative_words))

        total_sentiment_words = positive_count + negative_count

        if total_sentiment_words == 0:
            return {'neutral': 1.0}

        return {
            'positive': positive_count / total_sentiment_words,
            'negative': negative_count / total_sentiment_words,
            'neutral': 0.0
        }