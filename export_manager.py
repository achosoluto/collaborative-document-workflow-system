"""
Export System for Insights and Summaries
Supports multiple export formats (PDF, JSON, CSV, Excel, Word)
"""

import os
import json
import csv
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path

# Export libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Custom imports
from .summarization_engine import SummarizationResult
from .insight_extractor import ExtractedInsight
from .content_categorizer import Category, Tag
from .insight_dashboard import DashboardGenerator

logger = logging.getLogger(__name__)


class ExportFormat:
    """Supported export formats"""
    PDF = "pdf"
    JSON = "json"
    CSV = "csv"
    EXCEL = "excel"
    WORD = "word"
    HTML = "html"


class ExportManager:
    """Main export manager for insights and summaries"""

    def __init__(self):
        self.output_dir = Path("exports")
        self.output_dir.mkdir(exist_ok=True)

        # Initialize styles for PDF export
        self._initialize_pdf_styles()

    def _initialize_pdf_styles(self):
        """Initialize PDF styling"""
        self.styles = getSampleStyleSheet()

        # Custom styles
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=18,
            spaceAfter=20,
            textColor=colors.darkblue
        ))

        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkgreen
        ))

        self.styles.add(ParagraphStyle(
            name='InsightContent',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=8,
            leftIndent=20
        ))

    def export_summary(self, summary_result: SummarizationResult, format_type: str,
                      output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Export a summarization result

        Args:
            summary_result: Result from summarization engine
            format_type: Export format (pdf, json, csv, etc.)
            output_path: Custom output path (optional)

        Returns:
            Export result with file path and metadata
        """
        try:
            # Generate filename if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"summary_{timestamp}.{format_type}"
                output_path = self.output_dir / filename

            # Export based on format
            if format_type.lower() == ExportFormat.PDF:
                return self._export_summary_pdf(summary_result, output_path)
            elif format_type.lower() == ExportFormat.JSON:
                return self._export_summary_json(summary_result, output_path)
            elif format_type.lower() == ExportFormat.CSV:
                return self._export_summary_csv(summary_result, output_path)
            elif format_type.lower() == ExportFormat.WORD:
                return self._export_summary_word(summary_result, output_path)
            elif format_type.lower() == ExportFormat.HTML:
                return self._export_summary_html(summary_result, output_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported format: {format_type}'
                }

        except Exception as e:
            logger.error(f"Summary export error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def export_insights(self, insights: List[ExtractedInsight], format_type: str,
                       output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Export extracted insights

        Args:
            insights: List of extracted insights
            format_type: Export format
            output_path: Custom output path (optional)

        Returns:
            Export result with file path and metadata
        """
        try:
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"insights_{timestamp}.{format_type}"
                output_path = self.output_dir / filename

            if format_type.lower() == ExportFormat.PDF:
                return self._export_insights_pdf(insights, output_path)
            elif format_type.lower() == ExportFormat.JSON:
                return self._export_insights_json(insights, output_path)
            elif format_type.lower() == ExportFormat.CSV:
                return self._export_insights_csv(insights, output_path)
            elif format_type.lower() == ExportFormat.EXCEL:
                return self._export_insights_excel(insights, output_path)
            elif format_type.lower() == ExportFormat.WORD:
                return self._export_insights_word(insights, output_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported format: {format_type}'
                }

        except Exception as e:
            logger.error(f"Insights export error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def export_dashboard(self, dashboard_data: Dict[str, Any], format_type: str,
                        output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Export dashboard data

        Args:
            dashboard_data: Dashboard data from dashboard generator
            format_type: Export format
            output_path: Custom output path (optional)

        Returns:
            Export result with file path and metadata
        """
        try:
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"dashboard_{timestamp}.{format_type}"
                output_path = self.output_dir / filename

            if format_type.lower() == ExportFormat.PDF:
                return self._export_dashboard_pdf(dashboard_data, output_path)
            elif format_type.lower() == ExportFormat.JSON:
                return self._export_dashboard_json(dashboard_data, output_path)
            elif format_type.lower() == ExportFormat.HTML:
                return self._export_dashboard_html(dashboard_data, output_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported format for dashboard: {format_type}'
                }

        except Exception as e:
            logger.error(f"Dashboard export error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def _export_summary_pdf(self, summary_result: SummarizationResult, output_path: str) -> Dict[str, Any]:
        """Export summary as PDF"""
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        story = []

        # Title
        story.append(Paragraph("Document Summary Report", self.styles['CustomTitle']))
        story.append(Spacer(1, 12))

        # Summary metadata
        metadata_text = f"""
        <b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
        <b>Method:</b> {summary_result.method_used}<br/>
        <b>Original Length:</b> {summary_result.original_length} characters<br/>
        <b>Summary Length:</b> {summary_result.summary_length} characters<br/>
        <b>Compression Ratio:</b> {summary_result.compression_ratio".2%"}<br/>
        <b>Processing Time:</b> {summary_result.processing_time".2f"} seconds<br/>
        """

        story.append(Paragraph(metadata_text, self.styles['Normal']))
        story.append(Spacer(1, 12))

        # Summary content
        story.append(Paragraph("Summary:", self.styles['SectionHeader']))

        # Split summary into paragraphs for better formatting
        summary_paragraphs = summary_result.summary.split('\n\n')
        for para in summary_paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), self.styles['Normal']))
                story.append(Spacer(1, 8))

        # Key phrases section
        if summary_result.key_phrases:
            story.append(Spacer(1, 12))
            story.append(Paragraph("Key Phrases:", self.styles['SectionHeader']))

            phrases_text = ", ".join(summary_result.key_phrases[:10])  # Top 10 phrases
            story.append(Paragraph(phrases_text, self.styles['Normal']))

        # Build PDF
        doc.build(story)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'pdf',
            'file_size': os.path.getsize(output_path),
            'exported_sections': ['summary', 'metadata', 'key_phrases']
        }

    def _export_summary_json(self, summary_result: SummarizationResult, output_path: str) -> Dict[str, Any]:
        """Export summary as JSON"""
        export_data = {
            'export_info': {
                'exported_at': datetime.now().isoformat(),
                'export_format': 'json',
                'type': 'summary'
            },
            'summary_data': summary_result.to_dict()
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'json',
            'file_size': os.path.getsize(output_path)
        }

    def _export_summary_csv(self, summary_result: SummarizationResult, output_path: str) -> Dict[str, Any]:
        """Export summary as CSV"""
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)

            # Write metadata
            writer.writerow(['Field', 'Value'])
            writer.writerow(['Method Used', summary_result.method_used])
            writer.writerow(['Original Length', summary_result.original_length])
            writer.writerow(['Summary Length', summary_result.summary_length])
            writer.writerow(['Compression Ratio', f"{summary_result.compression_ratio".2%"}"])
            writer.writerow(['Processing Time', f"{summary_result.processing_time".2f"} seconds"])
            writer.writerow([])

            # Write summary content
            writer.writerow(['Summary Content'])
            # Split summary into chunks for CSV
            summary_lines = summary_result.summary.split('\n')
            for line in summary_lines:
                writer.writerow([line.strip()])

            writer.writerow([])
            writer.writerow(['Key Phrases'])
            for phrase in summary_result.key_phrases:
                writer.writerow([phrase])

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'csv',
            'file_size': os.path.getsize(output_path)
        }

    def _export_summary_word(self, summary_result: SummarizationResult, output_path: str) -> Dict[str, Any]:
        """Export summary as Word document"""
        doc = Document()
        doc.add_heading('Document Summary Report', 0)

        # Add metadata
        doc.add_heading('Summary Information', level=1)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Method: {summary_result.method_used}')
        doc.add_paragraph(f'Original Length: {summary_result.original_length} characters')
        doc.add_paragraph(f'Summary Length: {summary_result.summary_length} characters')
        doc.add_paragraph(f'Compression Ratio: {summary_result.compression_ratio".2%"}')
        doc.add_paragraph(f'Processing Time: {summary_result.processing_time".2f"} seconds')

        # Add summary content
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary_result.summary)

        # Add key phrases
        if summary_result.key_phrases:
            doc.add_heading('Key Phrases', level=1)
            for phrase in summary_result.key_phrases[:10]:
                doc.add_paragraph(f'• {phrase}', style='List Bullet')

        doc.save(str(output_path))

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'word',
            'file_size': os.path.getsize(output_path)
        }

    def _export_summary_html(self, summary_result: SummarizationResult, output_path: str) -> Dict[str, Any]:
        """Export summary as HTML"""
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Document Summary Report</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .header {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                .metadata {{ background-color: #f8f9fa; padding: 15px; margin: 20px 0; border-radius: 5px; }}
                .summary {{ line-height: 1.6; margin: 20px 0; }}
                .key-phrases {{ background-color: #e8f5e8; padding: 15px; margin: 20px 0; border-radius: 5px; }}
                .metric {{ display: inline-block; margin: 10px; padding: 10px; background-color: #fff; border: 1px solid #ddd; border-radius: 3px; }}
            </style>
        </head>
        <body>
            <h1 class="header">Document Summary Report</h1>

            <div class="metadata">
                <h2>Summary Information</h2>
                <div class="metric"><strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
                <div class="metric"><strong>Method:</strong> {summary_result.method_used}</div>
                <div class="metric"><strong>Original Length:</strong> {summary_result.original_length} characters</div>
                <div class="metric"><strong>Summary Length:</strong> {summary_result.summary_length} characters</div>
                <div class="metric"><strong>Compression Ratio:</strong> {summary_result.compression_ratio".2%"}</div>
                <div class="metric"><strong>Processing Time:</strong> {summary_result.processing_time".2f"} seconds</div>
            </div>

            <h2>Summary</h2>
            <div class="summary">{summary_result.summary.replace(chr(10), '<br>')}</div>

            <h2>Key Phrases</h2>
            <div class="key-phrases">
        """

        if summary_result.key_phrases:
            for phrase in summary_result.key_phrases[:10]:
                html_content += f'<div>• {phrase}</div>'

        html_content += """
            </div>
        </body>
        </html>
        """

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'html',
            'file_size': os.path.getsize(output_path)
        }

    def _export_insights_pdf(self, insights: List[ExtractedInsight], output_path: str) -> Dict[str, Any]:
        """Export insights as PDF"""
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        story = []

        # Title
        story.append(Paragraph("Insights Extraction Report", self.styles['CustomTitle']))
        story.append(Spacer(1, 12))

        # Summary statistics
        total_insights = len(insights)
        confidence_counts = {}
        for insight in insights:
            conf = insight.confidence.value
            confidence_counts[conf] = confidence_counts.get(conf, 0) + 1

        stats_text = f"""
        <b>Total Insights:</b> {total_insights}<br/>
        <b>High Confidence:</b> {confidence_counts.get('high', 0)}<br/>
        <b>Medium Confidence:</b> {confidence_counts.get('medium', 0)}<br/>
        <b>Low Confidence:</b> {confidence_counts.get('low', 0)}<br/>
        """

        story.append(Paragraph(stats_text, self.styles['Normal']))
        story.append(Spacer(1, 12))

        # Group insights by type
        insights_by_type = {}
        for insight in insights:
            insight_type = insight.insight_type.value
            if insight_type not in insights_by_type:
                insights_by_type[insight_type] = []
            insights_by_type[insight_type].append(insight)

        # Add each insight type section
        for insight_type, type_insights in insights_by_type.items():
            story.append(Paragraph(f"{insight_type.title()} ({len(type_insights)})", self.styles['SectionHeader']))

            for insight in type_insights[:20]:  # Limit to 20 per type for PDF
                content = f"[{insight.confidence.value.upper()}] {insight.content}"
                story.append(Paragraph(content, self.styles['InsightContent']))

                if insight.context:
                    context = f"Context: {insight.context[:100]}..."
                    story.append(Paragraph(context, self.styles['Normal']))
                    story.append(Spacer(1, 4))

            story.append(Spacer(1, 8))

        doc.build(story)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'pdf',
            'file_size': os.path.getsize(output_path),
            'insights_exported': total_insights
        }

    def _export_insights_json(self, insights: List[ExtractedInsight], output_path: str) -> Dict[str, Any]:
        """Export insights as JSON"""
        export_data = {
            'export_info': {
                'exported_at': datetime.now().isoformat(),
                'export_format': 'json',
                'type': 'insights',
                'total_insights': len(insights)
            },
            'insights': [insight.to_dict() for insight in insights]
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'json',
            'file_size': os.path.getsize(output_path)
        }

    def _export_insights_csv(self, insights: List[ExtractedInsight], output_path: str) -> Dict[str, Any]:
        """Export insights as CSV"""
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)

            # Write header
            writer.writerow(['Type', 'Content', 'Confidence', 'Context', 'Location'])

            # Write insights
            for insight in insights:
                writer.writerow([
                    insight.insight_type.value,
                    insight.content.replace('\n', ' ').replace('\r', ' ')[:500],  # Limit length
                    insight.confidence.value,
                    insight.context.replace('\n', ' ').replace('\r', ' ')[:200],  # Limit length
                    str(insight.location)
                ])

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'csv',
            'file_size': os.path.getsize(output_path)
        }

    def _export_insights_excel(self, insights: List[ExtractedInsight], output_path: str) -> Dict[str, Any]:
        """Export insights as Excel"""
        # Prepare data for DataFrame
        data = []
        for insight in insights:
            data.append({
                'Type': insight.insight_type.value,
                'Content': insight.content,
                'Confidence': insight.confidence.value,
                'Context': insight.context,
                'Location': json.dumps(insight.location),
                'Metadata': json.dumps(insight.metadata)
            })

        df = pd.DataFrame(data)

        # Create Excel writer with multiple sheets
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Main insights sheet
            df.to_excel(writer, sheet_name='Insights', index=False)

            # Summary sheet
            summary_data = []
            for insight_type in set(insight.insight_type.value for insight in insights):
                type_insights = [i for i in insights if i.insight_type.value == insight_type]
                summary_data.append({
                    'Type': insight_type,
                    'Count': len(type_insights),
                    'High_Confidence': len([i for i in type_insights if i.confidence.value == 'high']),
                    'Medium_Confidence': len([i for i in type_insights if i.confidence.value == 'medium']),
                    'Low_Confidence': len([i for i in type_insights if i.confidence.value == 'low'])
                })

            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'excel',
            'file_size': os.path.getsize(output_path)
        }

    def _export_insights_word(self, insights: List[ExtractedInsight], output_path: str) -> Dict[str, Any]:
        """Export insights as Word document"""
        doc = Document()
        doc.add_heading('Insights Extraction Report', 0)

        # Add summary
        doc.add_heading('Summary', level=1)
        total_insights = len(insights)

        confidence_counts = {}
        for insight in insights:
            conf = insight.confidence.value
            confidence_counts[conf] = confidence_counts.get(conf, 0) + 1

        doc.add_paragraph(f'Total Insights: {total_insights}')
        doc.add_paragraph(f'High Confidence: {confidence_counts.get("high", 0)}')
        doc.add_paragraph(f'Medium Confidence: {confidence_counts.get("medium", 0)}')
        doc.add_paragraph(f'Low Confidence: {confidence_counts.get("low", 0)}')

        # Group by type
        insights_by_type = {}
        for insight in insights:
            insight_type = insight.insight_type.value
            if insight_type not in insights_by_type:
                insights_by_type[insight_type] = []
            insights_by_type[insight_type].append(insight)

        # Add each type section
        for insight_type, type_insights in insights_by_type.items():
            doc.add_heading(f'{insight_type.title()} ({len(type_insights)})', level=1)

            for insight in type_insights[:50]:  # Limit for Word document
                # Add insight content
                para = doc.add_paragraph()
                para.add_run(f'[{insight.confidence.value.upper()}] ').bold = True
                para.add_run(insight.content)

                # Add context if available
                if insight.context:
                    doc.add_paragraph(f'Context: {insight.context[:200]}...')

        doc.save(str(output_path))

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'word',
            'file_size': os.path.getsize(output_path)
        }

    def _export_dashboard_pdf(self, dashboard_data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
        """Export dashboard as PDF"""
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        story = []

        # Title
        story.append(Paragraph("Document Insights Dashboard", self.styles['CustomTitle']))
        story.append(Spacer(1, 12))

        # Dashboard summary
        summary = dashboard_data.get('summary', {})
        if summary:
            story.append(Paragraph("Dashboard Summary", self.styles['SectionHeader']))

            summary_text = f"""
            <b>Total Insights:</b> {summary.get('total_insights', 0)}<br/>
            <b>Total Relationships:</b> {summary.get('total_relationships', 0)}<br/>
            <b>Total Gaps:</b> {summary.get('total_gaps', 0)}<br/>
            <b>Health Score:</b> {summary.get('health_score', 0)}/100<br/>
            """

            story.append(Paragraph(summary_text, self.styles['Normal']))
            story.append(Spacer(1, 12))

        # Key insights
        key_insights = summary.get('key_insights', [])
        if key_insights:
            story.append(Paragraph("Key Insights", self.styles['SectionHeader']))
            for insight in key_insights:
                story.append(Paragraph(f"• {insight}", self.styles['Normal']))
            story.append(Spacer(1, 12))

        # Recommendations
        recommendations = summary.get('recommendations', [])
        if recommendations:
            story.append(Paragraph("Recommendations", self.styles['SectionHeader']))
            for rec in recommendations:
                story.append(Paragraph(f"• {rec}", self.styles['Normal']))

        doc.build(story)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'pdf',
            'file_size': os.path.getsize(output_path)
        }

    def _export_dashboard_json(self, dashboard_data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
        """Export dashboard as JSON"""
        export_data = {
            'export_info': {
                'exported_at': datetime.now().isoformat(),
                'export_format': 'json',
                'type': 'dashboard'
            },
            'dashboard_data': dashboard_data
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'json',
            'file_size': os.path.getsize(output_path)
        }

    def _export_dashboard_html(self, dashboard_data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
        """Export dashboard as HTML"""
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Document Insights Dashboard</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .header {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                .section {{ margin: 30px 0; }}
                .metric {{ background-color: #f8f9fa; padding: 15px; margin: 10px 0; border-radius: 5px; }}
                .insight {{ background-color: #e8f5e8; padding: 10px; margin: 5px 0; border-radius: 3px; }}
                .recommendation {{ background-color: #fff3cd; padding: 10px; margin: 5px 0; border-radius: 3px; }}
            </style>
        </head>
        <body>
            <h1 class="header">Document Insights Dashboard</h1>

            <div class="section">
                <h2>Dashboard Summary</h2>
        """

        summary = dashboard_data.get('summary', {})
        if summary:
            html_content += f"""
                <div class="metric">
                    <strong>Total Insights:</strong> {summary.get('total_insights', 0)}<br/>
                    <strong>Total Relationships:</strong> {summary.get('total_relationships', 0)}<br/>
                    <strong>Total Gaps:</strong> {summary.get('total_gaps', 0)}<br/>
                    <strong>Health Score:</strong> {summary.get('health_score', 0)}/100
                </div>
            """

        # Key insights
        key_insights = summary.get('key_insights', [])
        if key_insights:
            html_content += "<h3>Key Insights</h3>"
            for insight in key_insights:
                html_content += f'<div class="insight">• {insight}</div>'

        # Recommendations
        recommendations = summary.get('recommendations', [])
        if recommendations:
            html_content += "<h3>Recommendations</h3>"
            for rec in recommendations:
                html_content += f'<div class="recommendation">• {rec}</div>'

        html_content += """
            </div>
        </body>
        </html>
        """

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return {
            'success': True,
            'file_path': str(output_path),
            'format': 'html',
            'file_size': os.path.getsize(output_path)
        }

    def batch_export(self, data_items: List[Dict[str, Any]], format_type: str,
                    output_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Export multiple items in batch

        Args:
            data_items: List of {'type': str, 'data': Any, 'filename': str} items
            format_type: Export format
            output_dir: Output directory (optional)

        Returns:
            Batch export results
        """
        if output_dir:
            output_path = Path(output_dir)
            output_path.mkdir(exist_ok=True)
        else:
            output_path = self.output_dir

        results = []
        errors = []

        for item in data_items:
            try:
                item_type = item.get('type')
                data = item.get('data')
                filename = item.get('filename', f'export_{len(results)}')

                if item_type == 'summary' and hasattr(data, 'to_dict'):
                    # Handle SummarizationResult
                    result = self.export_summary(data, format_type, output_path / filename)
                elif item_type == 'insights':
                    # Handle insights list
                    result = self.export_insights(data, format_type, output_path / filename)
                elif item_type == 'dashboard':
                    # Handle dashboard data
                    result = self.export_dashboard(data, format_type, output_path / filename)
                else:
                    errors.append(f"Unknown item type: {item_type}")
                    continue

                results.append(result)

            except Exception as e:
                errors.append(str(e))

        return {
            'success': len(errors) == 0,
            'total_items': len(data_items),
            'successful_exports': len(results),
            'errors': errors,
            'results': results
        }

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get supported export formats by data type"""
        return {
            'summary': ['pdf', 'json', 'csv', 'word', 'html'],
            'insights': ['pdf', 'json', 'csv', 'excel', 'word'],
            'dashboard': ['pdf', 'json', 'html']
        }


# Global export manager instance
export_manager = ExportManager()